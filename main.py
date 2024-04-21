# This is a sample Python script.
import os
import re
import string
import time

import PyPDF2
import boto3
import emoji
import pandas as pd
from docx import Document
from keras_preprocessing.text import text_to_word_sequence

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.

# pdf_folder_path = 'C:\\Users\\mi\\Downloads\\cv\\Senior'
s3 = boto3.resource('s3')
bucket_name = 'vp-cv'
# folder_path = 'Senior Project Manager'

folder_path = 'CV_DE'

bucket_name_result = 'cv-text'
# folder_path_result = 'Senior Project Manager'
folder_path_result = 'CV_DE'


def get_files_in_folder(bucket_name, folder_path):
    bucket = s3.Bucket(bucket_name)
    files = [obj.key for obj in bucket.objects.filter(Prefix=folder_path) if not obj.key.endswith('/')]
    return files


def timed_function(func):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        print(f"Time to run the function:  {end_time - start_time} seconds")
        return result

    return wrapper


# def get_files_in_folder(folder_path):
#     # Use glob to match the pathname pattern
#     file_list = glob.glob(folder_path + "/*")
#
#     return file_list

def write_to_file(name, tokens):
    text = ' '.join(tokens)
    with open(name, 'w', encoding='utf-8') as f:
        f.write(text)


def write_to_file_s3(bucket_name, folder_path, file_name, tokens):
    text = ' '.join(tokens)
    s3 = boto3.resource('s3')
    file_path_in_bucket = folder_path + "/" + file_name
    s3.Object(bucket_name, file_path_in_bucket).put(Body=text)


# @timed_function
def read_file(file_path):
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            text = ''

            for page_number in range(num_pages):
                page = pdf_reader.pages[page_number]
                text += page.extract_text()

            return text

    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        return text

    elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
        df = pd.read_excel(file_path)
        return df

    else:
        return "Unsupported file format"


def read_file_s3(bucket_name, file_path):
    s3_client = boto3.client('s3')
    s3_client.download_file(bucket_name, file_path, 'temp_file')

    if file_path.endswith('.pdf'):
        with open('temp_file', 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            text = ''

            for page_number in range(num_pages):
                page = pdf_reader.pages[page_number]
                text += page.extract_text()

            return text

    elif file_path.endswith('.docx'):
        doc = Document('temp_file')
        text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        return text

    elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
        df = pd.read_excel('temp_file')
        return df

    else:
        return "Unsupported file format"


def clean_and_tokenize_text(text):
    # Lowercase the text
    text = text.lower()

    # Remove HTML tags
    text = re.sub(r"<.*?>", "", text)

    # Remove punctuation
    text = ''.join(c for c in text if c not in string.punctuation)

    # Replace emojis with their names surrounded by spaces, e.g. ":grinning_face:"
    text = emoji.demojize(text, delimiters=(" ", " "))

    # Thay thế teencode bằng text tương tứng

    # Tokenize the text
    tokens = text_to_word_sequence(text)

    # Remove stopwords
    # stop_words = set(stopwords.words('english'))
    # tokens = [w for w in tokens if not w in stop_words]

    return tokens


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    # list_file = get_files_in_folder(pdf_folder_path)
    # print(list_file)
    files = get_files_in_folder(bucket_name, folder_path)
    for file in files:
        file_name = os.path.split(file)[-1]
        file_name = file_name.replace(".pdf", "") + ".txt"
        tokens = clean_and_tokenize_text(read_file_s3(bucket_name, file))

        write_to_file_s3(bucket_name_result, folder_path_result, file_name, tokens)
        # print(read_file(file))
